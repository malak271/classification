from flask import Flask
from flask import render_template
from datetime import datetime
from flask import Flask, request 
import pickle
import pandas as pd
from flask_cors import CORS
import time
import os
import docx2txt
import docx
import PyPDF2
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from PyPDF2 import PdfReader 
# import shutil 
import shutil 
from datetime import datetime
import fitz
import flask_wtf
from flask_wtf import FlaskForm
from flask_wtf.file import FileField, FileRequired
from werkzeug.utils import secure_filename

# Importing flask module in the project is mandatory
# An object of Flask class is our WSGI application.
 
# Flask constructor takes the name of
# current module (__name__) as argument.
app = Flask(__name__,template_folder='template')


@app.route('/searchText', methods =["GET", "POST"])
def searchText():

    path = 'documents' 
    text =  request.form['search']
    files = os.listdir(path)
    
    for file_name in files:
        #abs_path = os.path.abspath(file_name)
        abs_path= 'documents/'+file_name 
        
        if os.path.isdir(abs_path):
            searchText(abs_path)

        if os.path.isfile(abs_path):  

            if(file_name.endswith('.pdf')):
                # return file_name
                search_pdf(file_name,text)

            elif(file_name.endswith('.docx')):
                doc = docx.Document(abs_path) #docx2txt.process(file_name)
                # doc.save(file_name)
                for paragraph in doc.paragraphs:
                    if text in paragraph.text:
                        x = paragraph.text.split(text)
                        paragraph.clear()
                        for i in range(len(x)-1):
                            paragraph.add_run(x[i])
                            paragraph.add_run(" "+text+" ") .font.highlight_color = WD_COLOR_INDEX.YELLOW
                            paragraph.add_run(x[i+1])
                            doc.save("output/"+file_name)  

    return render_template('home.html')


@app.route('/sort', methods =["GET", "POST"])
def sort():
    name_of_dir = 'documents/'
  
# Storing list of all files
# in the given directory in list_of_files
    list_of_files = filter( lambda x: os.path.isfile
                       (os.path.join(name_of_dir, x)),
                        os.listdir("documents/") 
                        )
    
  
# Sort list of file names by title 
    list_of_files = sorted( list_of_files,
                        key = lambda x: sort_key(x))
    

# Iterate over sorted list of file 
# names and print them along with title one by one 
    for name_of_file in list_of_files:
        path_of_file = os.path.join(os.path.abspath(os.path.dirname(name_of_dir)), name_of_file)

        if(name_of_file.endswith('.pdf')):#expression_if_true if condition else expression_if_false
            title_of_file  = PdfReader(path_of_file).metadata.title
            if title_of_file  is not None:
               title_of_file+=".pdf" 
        elif(name_of_file.endswith('.docx')):
            title_of_file = getMetaData(docx.Document(path_of_file))
            if title_of_file  is not None:
               title_of_file+=".docx" 

        print(title_of_file, ' -->', name_of_file)

        if(title_of_file is None):
          title_of_file=name_of_file
         
        new_path=os.path.join('output_sort', title_of_file)
        shutil.copyfile(path_of_file, new_path)

    return render_template('home.html')

@app.route("/predict", methods=['POST'])
def predictAll():
    path="documents"
    result=dict()
    # os.chdir(path)
    files = os.listdir(path)
    #print(files)
    for file_name in files:
        #print(file_name)
        abs_path =os.path.join(
            os.path.abspath(os.path.dirname("documents/")),
            file_name)

        if os.path.isdir(abs_path):
            predict(abs_path)
            
        if os.path.isfile(abs_path):         
            if(file_name.endswith('.pdf')):
                my_text=readPDF(abs_path)
                result[file_name] = predict(my_text) 
            elif(file_name.endswith('.docx')):
                my_text = docx2txt.process(abs_path)
                result[file_name] = predict(my_text) 

            # result.append((predict(my_text)))

    return result

def predict(test_article):

    start = time.time()
    # test_article=request.form['x']
    feature = None
    rf_model = None
    with open(os.path.join("models/",
            "text_classification_tfidf_vectorizer"),'rb') as data:

        feature = pickle.load(data)
    # with open("models/text_classification_rf_model", 'rb') as data: #text_classification_tfidf_vectorizer
    with open(os.path.join("models/",
            "text_classification_rf_model"),'rb') as data:
        rf_model = pickle.load(data)
    
    test_article = test_article.lower()
    test_frame = pd.DataFrame({"Text":[test_article]})
    test_feature = feature.transform(test_frame.Text).toarray()
    prediction = rf_model.predict(test_feature)

    id_to_category = {0:"business",1:"entertainment",2:"politics",3:"sport",4:"tech"}
    prediction_cat = prediction[0]
    prediction_cat = id_to_category[prediction_cat]

    print ("Hey this News article belongs to ||",prediction_cat," || Category")
    print("Time spent handling the request: %f" % (time.time() - start))

    return prediction_cat
 
def readPDF(file_name):
    # creating a pdf file object
    pdfFileObj = open(file_name, 'rb')
   
# creating a pdf reader object
    pdfReader = PyPDF2.PdfReader(pdfFileObj)
   
# printing number of pages in pdf file
    # print(pdfReader.numPages)
   # reader.getPage(pageNumber) is deprecated and was removed in PyPDF2 3.0.0. Use reader.pages[page_number] instead.
# creating a page object
    pageObj = pdfReader.pages[0]
   
# extracting text from page
    # print(pageObj.extract_text())
   
# closing the pdf file object
    pdfFileObj.close()

    return pageObj.extract_text()

# The route() function of the Flask class is a decorator,
# which tells the application which URL should call
# the associated function.
@app.route('/')
# ‘/’ URL is bound with hello_world() function.
def hello_world():
    return render_template("home.html")
 
def search_pdf(file_name,text):


    ### READ IN PDF
    doc = fitz.open("documents/"+file_name)

    for page in doc:
        ### SEARCH
        text = text
        text_instances = page.search_for(text)#searchFor

        ### HIGHLIGHT
    for inst in text_instances:
        highlight = page.add_highlight_annot(inst)
        doc.save("output/"+file_name,garbage=4, deflate=True, clean=True)
    #  return request.form['search'] 


def getMetaData(doc):
    prop = doc.core_properties
    return prop.title


def sort_key(file_name):
    name_of_dir = 'documents/'
    if(file_name.endswith('.pdf')):
       title=PdfReader(open(os.path.join(name_of_dir, file_name), 'rb')).metadata.title
    elif(file_name.endswith('.docx')):
       title=getMetaData(docx.Document(os.path.join(name_of_dir, file_name)))
    print(title)
    if(title is None):
       return "no title"
    return title


app.config['SECRET_KEY'] = 'supersecretkey'
app.config['UPLOAD_FOLDER'] = 'static/files'

class UploadFileForm(FlaskForm):
    file = FileField(validators=[FileRequired()])
    # submit = SubmitField("Upload File")

@app.route('/upload', methods=['GET',"POST"])
def upload():
    form = UploadFileForm()
    if form.is_submitted():

        file = form.file.data # First grab the file

        file.save(os.path.join(
            os.path.abspath(os.path.dirname("documents/")),
            secure_filename(file.filename))) # Then save the file

        # return "File has been uploaded."
    return render_template('home.html', form=form)


# main driver function
if __name__ == '__main__':
 
    # run() method of Flask class runs the application
    # on the local development server.
    app.run(debug = True)
